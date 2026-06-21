from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import date
from collections import defaultdict
import io

from db.queries import get_absences, get_all_students, get_all_disciplines


async def generate_full_report(date_from: date = None, date_to: date = None,
                               discipline_id: int = None) -> io.BytesIO:
    absences = await get_absences(discipline_id=discipline_id,
                                  date_from=date_from, date_to=date_to)
    students = await get_all_students()

    by_student = defaultdict(list)
    for row in absences:
        by_student[row["student_id"]].append(row)

    doc = Document()
    section = doc.sections[0]
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(1.5)

    title = doc.add_heading("Сводка пропусков", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    period_parts = []
    if date_from:
        period_parts.append(f"с {date_from.strftime('%d.%m.%Y')}")
    if date_to:
        period_parts.append(f"по {date_to.strftime('%d.%m.%Y')}")
    period_str = " ".join(period_parts) if period_parts else "за всё время"

    sub_p = doc.add_paragraph(f"Период: {period_str}")
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.runs[0].font.size      = Pt(11)
    sub_p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # 5 колонок: №, Студент, Дисциплина, Пара, Часов
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    hdr     = table.rows[0].cells
    headers = ["№", "Студент", "Дисциплина", "Пара", "Часов"]
    for i, text in enumerate(headers):
        hdr[i].text = text
        hdr[i].paragraphs[0].runs[0].bold      = True
        hdr[i].paragraphs[0].alignment         = WD_ALIGN_PARAGRAPH.CENTER
        hdr[i].vertical_alignment              = WD_ALIGN_VERTICAL.CENTER

    table.columns[0].width = Cm(1)
    table.columns[1].width = Cm(4.5)
    table.columns[2].width = Cm(5.5)
    table.columns[3].width = Cm(2.5)
    table.columns[4].width = Cm(2.5)

    n           = 1
    total_hours = 0

    for student in students:
        sid  = student["id"]
        rows = by_student.get(sid, [])
        if not rows:
            continue

        for row in sorted(rows, key=lambda r: r["date"]):
            lesson_num  = row.get("lesson_num")
            lesson_time = row.get("lesson_time", "")

            if lesson_num:
                pair_cell = f"{lesson_num}-я"
                if lesson_time:
                    pair_cell += f"\n{lesson_time}"
            else:
                pair_cell = "—"

            row_cells = table.add_row().cells
            row_cells[0].text = str(n)
            row_cells[1].text = f"{student['last_name']} {student['first_name']}"
            row_cells[2].text = row["discipline"]
            row_cells[3].text = pair_cell
            row_cells[4].text = str(row["hours"])

            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            n           += 1
            total_hours += row["hours"]

    total_row = table.add_row().cells
    total_row[0].merge(total_row[1]).merge(total_row[2]).merge(total_row[3])
    total_row[0].text = "Итого часов:"
    total_row[0].paragraphs[0].runs[0].bold  = True
    total_row[0].paragraphs[0].alignment     = WD_ALIGN_PARAGRAPH.RIGHT
    total_row[4].text = str(total_hours)
    total_row[4].paragraphs[0].runs[0].bold  = True
    total_row[4].paragraphs[0].alignment     = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    footer = doc.add_paragraph(f"Документ сформирован: {date.today().strftime('%d.%m.%Y')}")
    footer.alignment      = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size      = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


async def generate_student_report(student_id: int) -> io.BytesIO:
    absences = await get_absences(student_id=student_id)
    if not absences:
        return None

    student   = absences[0]
    full_name = f"{student['last_name']} {student['first_name']}"

    doc     = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Cm(2)
    section.left_margin  = Cm(2.5)
    section.right_margin = Cm(1.5)

    h = doc.add_heading(f"Пропуски: {full_name}", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 5 колонок: Дата, Дисциплина, Пара, Часов, Уважительная
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for i, text in enumerate(["Дата", "Дисциплина", "Пара", "Часов", "Уважит."]):
        hdr[i].text = text
        hdr[i].paragraphs[0].runs[0].bold  = True
        hdr[i].paragraphs[0].alignment     = WD_ALIGN_PARAGRAPH.CENTER

    total = 0
    for row in absences:
        lesson_num  = row.get("lesson_num")
        lesson_time = row.get("lesson_time", "")

        if lesson_num:
            pair_cell = f"{lesson_num}-я"
            if lesson_time:
                pair_cell += f"\n{lesson_time}"
        else:
            pair_cell = "—"

        r    = table.add_row().cells
        r[0].text = row["date"].strftime("%d.%m.%Y")
        r[1].text = row["discipline"]
        r[2].text = pair_cell
        r[3].text = str(row["hours"])
        r[4].text = "✓" if row["confirmed"] else "—"

        for c in r:
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        total += row["hours"]

    tr = table.add_row().cells
    tr[0].merge(tr[1]).merge(tr[2]).merge(tr[3])
    tr[0].text = "Итого:"
    tr[0].paragraphs[0].runs[0].bold  = True
    tr[0].paragraphs[0].alignment     = WD_ALIGN_PARAGRAPH.RIGHT
    tr[4].text = str(total)
    tr[4].paragraphs[0].runs[0].bold  = True
    tr[4].paragraphs[0].alignment     = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
